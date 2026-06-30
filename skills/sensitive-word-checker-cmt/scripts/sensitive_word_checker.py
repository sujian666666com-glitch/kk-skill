#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
敏感词检测与标注工具 v3（内置词库版）
支持文件类型: PDF, Word (.docx), Excel (.xlsx/.xls)
核心改进:
  - v3: 内置65个敏感词词库（9类），无需外部词库文件即可使用
  - v2: 仅对匹配的敏感词本身标黄，不再整段/整句标黄
"""

import argparse
import os
import sys
import re
import copy
from pathlib import Path
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============================================================
# 内置敏感词库 (9类65个词条，按长度降序排列，长词优先匹配)
# ============================================================
BUILTIN_WORDS = [
    # --- 5字词 ---
    "多双边机制",      # 高规格名号
    "企业专题会",      # 企业专题

    # --- 4字词 ---
    "驻地记者",        # 收费合同
    "讲者幻灯",        # 报社文本

    # --- 3字词 ---
    "自贸港",          # 高规格名号
    "分论坛",          # 会议名称
    "卫星会",          # 会议名称
    "报告会",          # 会议名称
    "研讨会",          # 会议名称
    "培训班",          # 会议名称
    "上市会",          # 会议名称
    "推介会",          # 会议名称
    "研讨班",          # 会议名称
    "产品名",          # 名称
    "商品名",          # 名称

    # --- 2字词 ---
    "中国",            # 高规格名号
    "中华",            # 高规格名号
    "国家",            # 高规格名号
    "全国",            # 高规格名号
    "国际",            # 高规格名号
    "世界",            # 高规格名号
    "全球",            # 高规格名号
    "环球",            # 高规格名号
    "亚洲",            # 高规格名号
    "中外",            # 高规格名号
    "海外",            # 高规格名号
    "中西",            # 高规格名号
    "峰会",            # 高规格名号
    "高端",            # 高规格名号
    "高峰",            # 高规格名号
    "巅峰",            # 高规格名号
    "论坛",            # 会议名称
    "讲座",            # 会议名称
    "讲坛",            # 会议名称
    "大会",            # 会议名称
    "年会",            # 会议名称
    "评比",            # 评奖活动
    "评选",            # 评奖活动
    "大赛",            # 评奖活动
    "海选",            # 评奖活动
    "奖励",            # 评奖活动
    "示范",            # 评奖活动
    "表彰",            # 评奖活动
    "标杆",            # 评奖活动
    "比拼",            # 评奖活动
    "达标",            # 评奖活动
    "植入",            # 评奖/收费合同
    "学分",            # 评奖/收费合同
    "案例",            # 评奖活动
    "茶歇",            # 茶歇宴会
    "宴会",            # 茶歇宴会
    "宴请",            # 茶歇宴会
    "软文",            # 收费合同
    "报道",            # 收费合同
    "硬广",            # 收费合同
    "采访",            # 收费合同
    "新闻",            # 收费合同
    "稿件",            # 收费合同
    "文章",            # 收费合同
    "字数",            # 收费合同
    "记者",            # 收费合同
    "授牌",            # 其他
    "飞检",            # 其他

    # --- 1字词 ---
    "赛",              # 评奖活动
    "奖",              # 评奖活动
]

def get_builtin_words():
    """返回内置敏感词列表（已按长度降序排列）"""
    return list(BUILTIN_WORDS)


def get_desktop_path():
    """获取桌面路径"""
    try:
        import ctypes
        import ctypes.wintypes
        buf = ctypes.create_unicode_buffer(ctypes.wintypes.MAX_PATH)
        ctypes.windll.shell32.SHGetFolderPathW(None, 0, None, 0, buf)
        desktop = Path(buf.value)
        if desktop.exists():
            return desktop
    except:
        pass
    return Path.home() / "Desktop"


# ============================================================
# 敏感词提取 - 支持多种格式
# ============================================================

def extract_sensitive_words(word_file):
    """
    从Word文档中提取敏感词
    支持：
    - 纯词条列表（每行一个或逗号分隔）
    - 分类禁用规则格式（如"【严禁】xxx"）
    - 表格形式
    """
    from docx import Document

    all_words = []       # 所有原始词条
    keywords = []        # 拆分后的短关键词（用于精确匹配）

    doc = Document(word_file)

    # ---- 提取段落内容 ----
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or len(text) < 2:
            continue
        # 提取分类条目中的具体词汇
        words_from_text = _parse_text_for_words(text)
        all_words.extend(words_from_text)

    # ---- 提取表格内容 ----
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text and len(text) >= 2:
                        words_from_text = _parse_text_for_words(text)
                        all_words.extend(words_from_text)

    # 去重合并
    combined = list(set(all_words + keywords))

    # 过滤：只保留有实际检测价值的词
    final = [w for w in combined if _is_valid_word(w)]
    # 按长度降序排列（长的优先匹配，避免短词先吞掉长词的一部分）
    final.sort(key=len, reverse=True)

    return final


def _parse_text_for_words(text):
    """从一段文本中解析出所有敏感词"""
    words = []
    keywords = []

    # 模式1：去除【严禁】/【注意】等标签前缀
    clean_text = re.sub(r'[【\[].*?[】\]]', '', text).strip()
    if not clean_text or len(clean_text) < 2:
        return []

    # 模式2：按顿号、斜杠、逗号、引号等分隔符拆分
    separators_pattern = r'[,，;；、/／""''「」【】《》（）\n\r\t]+'
    parts = re.split(separators_pattern, clean_text)

    for part in parts:
        part = part.strip().strip('"\'').strip()
        if not part or len(part) < 2:
            continue

        # 记录原始词条
        words.append(part)

        # 对较长的短语做二次拆分（如 "中国、中华、国家、全国" 在同一串中）
        sub_parts = re.split(r'[、,，/／]', part)
        for sp in sub_parts:
            sp = sp.strip()
            # 保留2-10个字符的有意义词汇
            if 2 <= len(sp) <= 10 and re.search(r'[\u4e00-\u9fff]', sp):
                keywords.append(sp)

    return words


def _is_valid_word(word):
    """判断是否为有效的敏感词"""
    if not word or len(word) < 2:
        return False
    # 排除纯标点、纯数字等无意义内容
    if re.match(r'^[\d\s\W]+$', word):
        return False
    return True


# ============================================================
# DOCX 精确标黄 - 核心改进部分
# ============================================================

def check_and_highlight_docx(input_file, sensitive_words, output_file):
    """
    检测并高亮Word文档中的敏感词（仅标黄匹配文字本身）
    """
    from docx import Document

    doc = Document(input_file)
    found_words = set()

    def process_paragraph_element(p_element):
        """处理一个段落的XML元素，对其中的所有run进行精确替换"""
        modified = False

        # 获取所有直接子run元素
        runs = p_element.findall(qn('w:r'))
        new_children = []

        for run in runs:
            t_elements = run.findall(qn('w:t'))
            if not t_elements:
                new_children.append(run)
                continue

            # 收集run中所有文本
            full_run_text = ''.join(t.text or '' for t in t_elements)
            if not full_run_text or len(full_run_text) == 0:
                new_children.append(run)
                continue

            # 查找所有敏感词在当前文本中的匹配位置
            all_matches = []
            for word in sensitive_words:
                pattern = re.compile(re.escape(word))
                for m in pattern.finditer(full_run_text):
                    all_matches.append((m.start(), m.end(), word))

            if not all_matches:
                new_children.append(run)
                continue

            # 按位置排序并去重叠（保留最长的匹配）
            all_matches.sort(key=lambda x: (x[0], -(x[1] - x[0])))
            merged_matches = _merge_overlaps(all_matches)

            if not merged_matches:
                new_children.append(run)
                continue

            # 有命中！记录发现的敏感词
            for _, _, w in merged_matches:
                found_words.add(w)

            # 复制原run的属性（字体、大小、加粗等）
            base_rpr = run.find(qn('w:rPr'))

            # 将文本按照匹配位置切分为片段
            segments = []
            last_end = 0
            for start, end, w in merged_matches:
                if start > last_end:
                    segments.append((full_run_text[last_end:start], False))
                segments.append((full_run_text[start:end], True))
                last_end = end
            if last_end < len(full_run_text):
                segments.append((full_run_text[last_end:], False))

            # 为每个片段创建新的run
            for seg_text, is_sens in segments:
                new_run = OxmlElement('w:r')

                # 复制原有格式
                if base_rpr is not None:
                    new_rpr = copy.deepcopy(base_rpr)
                    new_run.append(new_rpr)

                t_elem = OxmlElement('w:t')
                t_elem.text = seg_text
                t_elem.set(qn('xml:space'), 'preserve')
                new_run.append(t_elem)

                # 如果是敏感词片段，添加黄色高亮
                if is_sens:
                    rpr = new_run.find(qn('w:rPr'))
                    if rpr is None:
                        rpr = OxmlElement('w:rPr')
                        new_run.insert(0, rpr)

                    # 黄色底色
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'FFFF00')
                    shd.set(qn('w:val'), 'clear')
                    rpr.append(shd)

                    # 黄色字体高亮
                    hlt = OxmlElement('w:highlight')
                    hlt.set(qn('w:val'), 'yellow')
                    rpr.append(hlt)

                new_children.append(new_run)

            modified = True

        if modified:
            # 替换所有子元素
            for child in list(p_element):
                p_element.remove(child)
            for child in new_children:
                p_element.append(child)


    def _merge_overlaps(matches):
        """合并重叠的匹配区间，优先保留更长的匹配"""
        if not matches:
            return []
        merged = []
        for start, end, word in matches:
            if merged and start < merged[-1][1]:
                # 有重叠，保留更长的一个
                prev_start, prev_end, prev_word = merged[-1]
                if (end - start) > (prev_end - prev_start):
                    merged[-1] = (start, end, word)
                else:
                    # 延伸前一个区间的结束位置
                    if end > prev_end:
                        merged[-1] = (prev_start, end, prev_word)
            else:
                merged.append([start, end, word])
        return merged

    # 处理正文段落
    for paragraph in doc.paragraphs:
        process_paragraph_element(paragraph._p)

    # 处理表格中的段落
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph_element(paragraph._p)

    doc.save(output_file)
    return list(found_words)


# ============================================================
# Excel 高亮标注
# ============================================================

def check_and_highlight_xlsx(input_file, sensitive_words, output_file):
    """检测并高亮Excel文档中的敏感词"""
    import openpyxl
    from openpyxl.styles import PatternFill

    wb = openpyxl.load_workbook(input_file)
    found_words = set()
    yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value:
                    cell_text = str(cell.value)
                    for word in sensitive_words:
                        if re.search(re.escape(word), cell_text):
                            found_words.add(word)
                            cell.fill = yellow_fill

    wb.save(output_file)
    return list(found_words)


# ============================================================
# PDF 标注
# ============================================================

def check_and_highlight_pdf_pymupdf(input_file, sensitive_words, output_file):
    """检测并高亮PDF中的敏感词 (使用pymupdf)"""
    import fitz

    doc = fitz.open(input_file)
    found_words = set()

    for page_num in range(len(doc)):
        page = doc[page_num]

        for word in sensitive_words:
            pattern = re.compile(re.escape(word))
            text_instances = page.search_for(word)

            if text_instances:
                found_words.add(word)
                for inst in text_instances:
                    highlight = page.add_highlight_annot(inst)
                    highlight.set_colors(stroke=(1, 1, 0))
                    highlight.update()

    doc.save(output_file)
    return list(found_words)


def check_pdf_fallback(input_file, sensitive_words, output_file):
    """PDF检测回退方案 (使用PyPDF2，只能报告不能标注)"""
    from PyPDF2 import PdfReader
    import shutil

    found_words = set()
    reader = PdfReader(input_file)
    full_text = ""

    for page in reader.pages:
        try:
            text = page.extract_text()
            if text:
                full_text += text + "\n"
        except Exception as e:
            print(f"警告: 无法提取部分页面文本: {e}")

    for word in sensitive_words:
        if re.search(re.escape(word), full_text):
            found_words.add(word)

    os.makedirs(os.path.dirname(str(output_file)) if os.path.dirname(str(output_file)) else '.', exist_ok=True)
    shutil.copy(input_file, str(output_file))

    return list(found_words)


# ============================================================
# 主入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(description='敏感词检测与标注工具 v3（内置词库版）')
    parser.add_argument('--words', help='敏感词库文件路径 (.docx)，不指定则使用内置词库')
    parser.add_argument('--input', required=True, help='要检测的文件路径')
    parser.add_argument('--output', help='输出目录路径 (默认: 桌面)')

    args = parser.parse_args()

    # 检查输入文件存在性
    if not os.path.exists(args.input):
        print(f"错误: 输入文件不存在: {args.input}")
        sys.exit(1)

    # 提取敏感词：优先使用外部词库，否则使用内置词库
    if args.words:
        if not os.path.exists(args.words):
            print(f"错误: 敏感词库文件不存在: {args.words}")
            sys.exit(1)
        print(f"[1/3] 正在读取外部敏感词库: {args.words}")
        sensitive_words = extract_sensitive_words(args.words)
    else:
        print(f"[1/3] 使用内置敏感词库")
        sensitive_words = get_builtin_words()
    print(f"      已加载 {len(sensitive_words)} 个敏感词")

    if not sensitive_words:
        print("警告: 敏感词库为空!")

    # 输出路径
    input_path = Path(args.input)
    output_dir = Path(args.output) if args.output else get_desktop_path()
    output_file = output_dir / input_path.name

    ext = input_path.suffix.lower()
    print(f"[2/3] 正在检测文件: {args.input} ({ext})")

    # 检查pymupdf
    has_pymupdf = False
    try:
        import fitz
        has_pymupdf = True
    except ImportError:
        pass

    # 执行检测
    found_words = []
    try:
        if ext == '.docx':
            found_words = check_and_highlight_docx(args.input, sensitive_words, str(output_file))
        elif ext in ('.xlsx', '.xls'):
            found_words = check_and_highlight_xlsx(args.input, sensitive_words, str(output_file))
        elif ext == '.pdf':
            if has_pymupdf:
                found_words = check_and_highlight_pdf_pymupdf(args.input, sensitive_words, str(output_file))
            else:
                found_words = check_pdf_fallback(args.input, sensitive_words, str(output_file))
        else:
            print(f"错误: 不支持的文件类型: {ext}")
            print("支持的类型: .docx, .pdf, .xlsx, .xls")
            sys.exit(1)
    except Exception as e:
        print(f"处理文件时出错: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

    # 输出结果
    print(f"[3/3] 检测完成!")
    print(f"\n{'='*50}")
    print(f"输出文件: {output_file}")

    if found_words:
        print(f"\n发现 {len(found_words)} 个敏感词:")
        for i, word in enumerate(found_words, 1):
            print(f"  {i}. {word}")
    else:
        print("\n未发现敏感词")

    return found_words


if __name__ == '__main__':
    main()
